"""Parser fidelity: page/section/char spans on generated + Persian fixtures.

These exercise the real parser dependencies (pypdf / python-docx) so they are
skipped where those are not installed (they run in CI). They do NOT require a
database.
"""

from __future__ import annotations

import pytest

from app.ingestion.parsers.text import parse_text
from tests.integration.fixtures_docs import PERSIAN_TEXT, make_text_pdf


def test_pdf_per_page_offsets() -> None:
    pytest.importorskip("pypdf")
    from app.ingestion.parsers.pdf import parse_pdf

    pdf = make_text_pdf(["Hello page one", "Second page text"])
    segments = parse_pdf(pdf)

    assert [s.page_no for s in segments] == [1, 2]
    assert "Hello page one" in segments[0].text
    assert "Second page text" in segments[1].text
    # Char offsets are monotonic and non-overlapping across pages.
    assert segments[0].char_start == 0
    assert segments[1].char_start >= segments[0].char_end


def test_pdf_image_only_is_no_text() -> None:
    pytest.importorskip("pypdf")
    from app.ingestion.guards import GuardError
    from app.ingestion.parsers.pdf import parse_pdf
    from app.models.enums import DocumentErrorCode

    # A PDF page with no text content stream => NO_TEXT (no OCR in v1).
    empty = make_text_pdf([""])
    with pytest.raises(GuardError) as exc:
        parse_pdf(empty)
    assert exc.value.code is DocumentErrorCode.NO_TEXT


def test_persian_text_parser_offsets() -> None:
    segments = parse_text(PERSIAN_TEXT.encode("utf-8"))
    assert len(segments) == 1
    seg = segments[0]
    assert seg.char_start == 0
    assert seg.char_end == len(seg.text)
    assert "فارسی" in seg.text
    # ZWNJ-joined word survives decode.
    assert "می‌رود" in seg.text


def test_docx_section_path() -> None:
    pytest.importorskip("docx")
    import io

    from docx import Document as DocxDocument

    from app.ingestion.parsers.docx import parse_docx

    buf = io.BytesIO()
    doc = DocxDocument()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Body text under chapter one.")
    doc.add_heading("Subsection", level=2)
    doc.add_paragraph("Deeper body text.")
    doc.save(buf)

    segments = parse_docx(buf.getvalue())
    bodies = [s for s in segments if s.text.startswith("Body") or s.text.startswith("Deeper")]
    assert any(s.section_path == "Chapter One" for s in bodies)
    assert any(s.section_path == "Chapter One > Subsection" for s in bodies)
